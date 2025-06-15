"""
Contextual Word Document Processor for ViettelPay Knowledge Base

This processor implements Anthropic's Contextual Retrieval technique:
- Extracts content from Word documents including text sections and tables
- Uses LLM to generate contextual information for each chunk
- Enhances chunks with context before embedding
"""

import re
from typing import List, Dict, Any, Optional
from pathlib import Path

from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain.schema import Document
from markitdown import MarkItDown

# Import LLM client factory and configuration
from src.llm.llm_client import LLMClientFactory, BaseLLMClient
from src.utils.config import get_google_api_key


class ContextualWordProcessor:
    """
    Contextual processor for Word documents using LLM enhancement
    """

    def __init__(
        self, llm_client=None, llm_provider: str = "gemini", api_key: str = None
    ):
        """
        Initialize the processor with an LLM client for contextual enhancement

        Args:
            llm_client: Pre-initialized LLM client instance (optional)
            llm_provider: LLM provider to use ("gemini" or "openai")
            api_key: API key for the LLM provider (optional, will use config if not provided)
        """
        if llm_client:
            self.llm_client = llm_client
        else:
            # Initialize LLM client using factory
            try:
                if llm_provider == "gemini":
                    api_key = api_key or get_google_api_key()
                    self.llm_client = LLMClientFactory.create_client(
                        "gemini", api_key=api_key, model="gemini-2.0-flash"
                    )
                    print(
                        "✅ ContextualWordProcessor initialized with Gemini 2.0 Flash"
                    )
                elif llm_provider == "openai":
                    from src.utils.config import get_openai_api_key

                    api_key = api_key or get_openai_api_key()
                    self.llm_client = LLMClientFactory.create_client(
                        "openai", api_key=api_key, model="gpt-4o-mini"
                    )
                    print(
                        "✅ ContextualWordProcessor initialized with OpenAI GPT-4o-mini"
                    )
                else:
                    raise ValueError(f"Unsupported LLM provider: {llm_provider}")
            except Exception as e:
                print(f"⚠️ Failed to initialize LLM client: {e}")
                print("   Context generation will be disabled")
                self.llm_client = None

        self.md_converter = MarkItDown()
        self.current_section_hierarchy = []
        self.section_counter = 0

        # Vietnamese contextual prompt template
        self.contextual_prompt_template = """<document>
{WHOLE_DOCUMENT}
</document>

Đây là đoạn văn bản cần được đặt trong ngữ cảnh của toàn bộ tài liệu:
<chunk>
{CHUNK_CONTENT}
</chunk>

Hãy cung cấp ngữ cảnh và tóm tắt ngắn gọn để giúp định vị đoạn văn này trong toàn bộ tài liệu ViettelPay Pro, nhằm cải thiện khả năng tìm kiếm thông tin. Chỉ trả lời bằng ngữ cảnh ngắn gọn, không cần giải thích thêm."""

    def process_word_document(self, file_path: str) -> List[Document]:
        """
        Main method to process a Word document into contextualized LangChain Documents

        Args:
            file_path: Path to the Word document

        Returns:
            List of contextualized LangChain Document objects
        """
        print(
            f"[INFO] Processing Word document with contextual enhancement: {file_path}"
        )

        if not Path(file_path).exists():
            raise FileNotFoundError(f"Word document not found: {file_path}")

        # Get the full document content using markitdown for context generation
        try:
            result = self.md_converter.convert(file_path)
            whole_document_content = result.text_content
            print(
                f"[INFO] Extracted full document content: {len(whole_document_content)} characters"
            )
        except Exception as e:
            print(f"[WARNING] Could not extract full document with markitdown: {e}")
            whole_document_content = ""

        # Load the document for structured processing
        doc = DocxDocument(file_path)

        # Extract all content with structure
        all_documents = []
        current_section_content = []
        current_section_title = "Introduction"
        current_section_level = 0

        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith("}p"):
                paragraph = Paragraph(element, doc)

                # Check if it's a heading
                heading_info = self._extract_heading_info(paragraph)

                if heading_info:
                    # Save previous section if it has content
                    if current_section_content:
                        section_doc = self._create_contextual_section_document(
                            current_section_title,
                            current_section_content,
                            current_section_level,
                            file_path,
                            whole_document_content,
                        )
                        if section_doc:
                            all_documents.append(section_doc)

                    # Start new section
                    current_section_title = heading_info["title"]
                    current_section_level = heading_info["level"]
                    current_section_content = []
                    self._update_section_hierarchy(
                        heading_info["level"], current_section_title
                    )

                else:
                    # Regular paragraph content
                    paragraph_text = paragraph.text.strip()
                    if paragraph_text:
                        current_section_content.append(paragraph_text)

            # Check if it's a table
            elif element.tag.endswith("}tbl"):
                table = Table(element, doc)
                table_documents = self._process_contextual_table(
                    table, current_section_title, file_path, whole_document_content
                )
                all_documents.extend(table_documents)

        # Don't forget the last section
        if current_section_content:
            section_doc = self._create_contextual_section_document(
                current_section_title,
                current_section_content,
                current_section_level,
                file_path,
                whole_document_content,
            )
            if section_doc:
                all_documents.append(section_doc)

        print(
            f"[SUCCESS] Extracted {len(all_documents)} contextualized documents from Word file"
        )
        return all_documents

    def _generate_context(self, chunk_content: str, whole_document: str) -> str:
        """
        Generate contextual information for a chunk using LLM

        Args:
            chunk_content: The text chunk to contextualize
            whole_document: The full document content for context

        Returns:
            Generated context string
        """
        if not self.llm_client or not whole_document:
            return ""

        try:
            # Prepare the prompt
            prompt = self.contextual_prompt_template.format(
                WHOLE_DOCUMENT=whole_document, CHUNK_CONTENT=chunk_content
            )

            # Call LLM API using the generic client interface
            context = self.llm_client.generate(
                prompt,
                max_output_tokens=200,  # For Gemini
                max_tokens=200,  # For OpenAI (will be ignored by Gemini)
                temperature=0.1,  # Low temperature for consistent context generation
            )

            return context.strip() if context else ""

        except Exception as e:
            print(f"[WARNING] Failed to generate context: {e}")
            return ""

    def _extract_heading_info(self, paragraph: Paragraph) -> Optional[Dict[str, Any]]:
        """
        Extract heading information from a paragraph
        """
        if paragraph.style.name.startswith("Heading"):
            try:
                level = int(paragraph.style.name.split()[-1])
                title = paragraph.text.strip()
                return {"title": title, "level": level}
            except (ValueError, IndexError):
                pass

        # Also check for manual heading patterns (like "# Title")
        text = paragraph.text.strip()
        if text.startswith("#"):
            level = len(text) - len(text.lstrip("#"))
            title = text.lstrip("#").strip()
            return {"title": title, "level": level}

        # Check for numbered sections like "1. Title", "1.1. Title"
        section_pattern = r"^(\d+\.(?:\d+\.)*)\s*(.+)$"
        match = re.match(section_pattern, text)
        if match:
            section_num = match.group(1)
            title = match.group(2)
            level = section_num.count(".")
            return {"title": title, "level": level}

        return None

    def _update_section_hierarchy(self, level: int, title: str):
        """Update the current section hierarchy"""
        # Trim hierarchy to current level
        self.current_section_hierarchy = self.current_section_hierarchy[: level - 1]

        # Add current section
        if len(self.current_section_hierarchy) < level:
            self.current_section_hierarchy.extend(
                [""] * (level - len(self.current_section_hierarchy))
            )

        if level <= len(self.current_section_hierarchy):
            self.current_section_hierarchy[level - 1] = title

    def _create_contextual_section_document(
        self,
        title: str,
        content: List[str],
        level: int,
        source_file: str,
        whole_document: str,
    ) -> Optional[Document]:
        """
        Create a contextualized Document object from section content
        """
        if not content:
            return None

        # Combine all paragraphs in the section
        original_content = f"# {title}\n\n" + "\n\n".join(content)

        # Generate contextual information
        context = self._generate_context(original_content, whole_document)

        # Combine context with original content
        if context:
            contextualized_content = f"{context}\n\n{original_content}"
            print(
                f"[INFO] Generated context for section '{title}': {len(context)} chars"
            )
        else:
            contextualized_content = original_content

        # Create metadata
        metadata = {
            "doc_type": "section",
            "section_title": title,
            "section_level": level,
            "section_hierarchy": " > ".join(
                filter(None, self.current_section_hierarchy)
            ),
            "source_file": Path(source_file).name,
            "content_type": "text_section",
            "section_id": f"section_{self.section_counter}",
            "has_context": bool(context),
            "original_content": original_content,
        }

        self.section_counter += 1

        return Document(page_content=contextualized_content, metadata=metadata)

    def _process_contextual_table(
        self, table: Table, current_section: str, source_file: str, whole_document: str
    ) -> List[Document]:
        """
        Process a table into multiple contextualized Document objects (one per row)
        """
        documents = []

        if not table.rows:
            return documents

        # Extract headers from first row
        headers = []
        first_row = table.rows[0]
        for cell in first_row.cells:
            headers.append(cell.text.strip())

        # Process each data row (skip header row)
        for row_idx, row in enumerate(table.rows[1:], 1):
            row_content = self._process_table_row(row, headers, row_idx)

            if row_content:
                # Generate contextual information for the table row
                context = self._generate_context(row_content, whole_document)

                # Combine context with original content
                if context:
                    contextualized_content = f"{context}\n\n{row_content}"
                    print(
                        f"[INFO] Generated context for table row {row_idx}: {len(context)} chars"
                    )
                else:
                    contextualized_content = row_content

                # Create metadata for the table row
                metadata = {
                    "doc_type": "table_row",
                    "section_title": current_section,
                    "section_hierarchy": " > ".join(
                        filter(None, self.current_section_hierarchy)
                    ),
                    "source_file": Path(source_file).name,
                    "content_type": "table_data",
                    "table_headers": " | ".join(headers),
                    "row_number": row_idx,
                    "table_id": f"table_{current_section}_{row_idx}",
                    "has_context": bool(context),
                    "original_content": row_content,
                }

                doc = Document(page_content=contextualized_content, metadata=metadata)
                documents.append(doc)

        return documents

    def _process_table_row(
        self, row, headers: List[str], row_idx: int
    ) -> Optional[str]:
        """
        Process a single table row into content string
        """
        row_data = []

        for cell in row.cells:
            cell_text = cell.text.strip()
            row_data.append(cell_text)

        # Skip empty rows
        if not any(row_data):
            return None

        # Create structured content from the row
        content_parts = []

        for header, cell_value in zip(headers, row_data):
            if cell_value:  # Only include non-empty cells
                content_parts.append(f"{header}: {cell_value}")

        if not content_parts:
            return None

        # Create the final content
        row_content = f"Bảng dữ liệu - Hàng {row_idx}:\n" + "\n".join(content_parts)

        return row_content

    def get_document_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """
        Get statistics about the processed documents
        """
        stats = {
            "total_documents": len(documents),
            "sections": 0,
            "table_rows": 0,
            "doc_types": {},
            "sections_by_level": {},
            "contextualized_docs": 0,
            "non_contextualized_docs": 0,
        }

        for doc in documents:
            doc_type = doc.metadata.get("doc_type", "unknown")
            stats["doc_types"][doc_type] = stats["doc_types"].get(doc_type, 0) + 1

            # Count contextualized vs non-contextualized
            if doc.metadata.get("has_context", False):
                stats["contextualized_docs"] += 1
            else:
                stats["non_contextualized_docs"] += 1

            if doc_type == "section":
                stats["sections"] += 1
                level = doc.metadata.get("section_level", 0)
                stats["sections_by_level"][level] = (
                    stats["sections_by_level"].get(level, 0) + 1
                )
            elif doc_type == "table_row":
                stats["table_rows"] += 1

        return stats


# Example usage and testing
if __name__ == "__main__":
    # Test the processor with Gemini 2.0 Flash (default)
    processor = ContextualWordProcessor(llm_provider="gemini")

    # Example file path (adjust as needed)
    test_file = "viettelpay_docs/raw/Nghiệp vụ.docx"

    try:
        documents = processor.process_word_document(test_file)

        # Show some example documents
        print(f"\n[INFO] First 3 documents:")
        for i, doc in enumerate(documents[:3]):
            print(f"\nDocument {i+1}:")
            print(f"Type: {doc.metadata.get('doc_type')}")
            print(f"Section: {doc.metadata.get('section_title')}")
            print(f"Has Context: {doc.metadata.get('has_context')}")
            print(f"Content preview: {doc.page_content[:200]}...")

        # Show stats
        stats = processor.get_document_stats(documents)
        print(f"\n[INFO] Processing statistics:")
        for key, value in stats.items():
            print(f"  {key}: {value}")

    except FileNotFoundError:
        print("[ERROR] Test file not found. Please adjust the file path.")
    except Exception as e:
        print(f"[ERROR] Error processing document: {e}")
