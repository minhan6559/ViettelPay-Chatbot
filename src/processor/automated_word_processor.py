"""
Automated Word Document Processor for ViettelPay Knowledge Base

This processor automatically extracts content from Word documents including:
- Text sections based on document structure
- Tables processed row by row
- Proper metadata for each document chunk
"""

import re
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path

from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain.schema import Document


class AutomatedWordProcessor:
    """
    Automated processor for Word documents that extracts both text content and tables
    """

    def __init__(self):
        self.current_section_hierarchy = []
        self.section_counter = 0

    def process_word_document(self, file_path: str) -> List[Document]:
        """
        Main method to process a Word document into LangChain Documents

        Args:
            file_path: Path to the Word document

        Returns:
            List of LangChain Document objects
        """
        print(f"[INFO] Processing Word document: {file_path}")

        if not Path(file_path).exists():
            raise FileNotFoundError(f"Word document not found: {file_path}")

        # Load the document
        doc = DocxDocument(file_path)

        # Extract all content with structure
        all_documents = []
        current_section_content = []
        current_section_title = "Introduction"
        current_section_level = 0

        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith("}p"):
                paragraph = Paragraph(element, doc)

                # Check if it's a heading
                heading_info = self._extract_heading_info(paragraph)

                if heading_info:
                    # Save previous section if it has content
                    if current_section_content:
                        section_doc = self._create_section_document(
                            current_section_title,
                            current_section_content,
                            current_section_level,
                            file_path,
                        )
                        if section_doc:
                            all_documents.append(section_doc)

                    # Start new section
                    current_section_title = heading_info["title"]
                    current_section_level = heading_info["level"]
                    current_section_content = []
                    self._update_section_hierarchy(
                        heading_info["level"], current_section_title
                    )

                else:
                    # Regular paragraph content
                    paragraph_text = paragraph.text.strip()
                    if paragraph_text:
                        current_section_content.append(paragraph_text)

            # Check if it's a table
            elif element.tag.endswith("}tbl"):
                table = Table(element, doc)
                table_documents = self._process_table(
                    table, current_section_title, file_path
                )
                all_documents.extend(table_documents)

        # Don't forget the last section
        if current_section_content:
            section_doc = self._create_section_document(
                current_section_title,
                current_section_content,
                current_section_level,
                file_path,
            )
            if section_doc:
                all_documents.append(section_doc)

        print(f"[SUCCESS] Extracted {len(all_documents)} documents from Word file")
        return all_documents

    def _extract_heading_info(self, paragraph: Paragraph) -> Optional[Dict[str, Any]]:
        """
        Extract heading information from a paragraph

        Returns:
            Dict with 'title' and 'level' if it's a heading, None otherwise
        """
        if paragraph.style.name.startswith("Heading"):
            try:
                level = int(paragraph.style.name.split()[-1])
                title = paragraph.text.strip()
                return {"title": title, "level": level}
            except (ValueError, IndexError):
                pass

        # Also check for manual heading patterns (like "# Title")
        text = paragraph.text.strip()
        if text.startswith("#"):
            level = len(text) - len(text.lstrip("#"))
            title = text.lstrip("#").strip()
            return {"title": title, "level": level}

        # Check for numbered sections like "1. Title", "1.1. Title"
        section_pattern = r"^(\d+\.(?:\d+\.)*)\s*(.+)$"
        match = re.match(section_pattern, text)
        if match:
            section_num = match.group(1)
            title = match.group(2)
            level = section_num.count(".")
            return {"title": title, "level": level}

        return None

    def _update_section_hierarchy(self, level: int, title: str):
        """Update the current section hierarchy"""
        # Trim hierarchy to current level
        self.current_section_hierarchy = self.current_section_hierarchy[: level - 1]

        # Add current section
        if len(self.current_section_hierarchy) < level:
            self.current_section_hierarchy.extend(
                [""] * (level - len(self.current_section_hierarchy))
            )

        if level <= len(self.current_section_hierarchy):
            self.current_section_hierarchy[level - 1] = title

    def _create_section_document(
        self, title: str, content: List[str], level: int, source_file: str
    ) -> Optional[Document]:
        """
        Create a Document object from section content
        """
        if not content:
            return None

        # Combine all paragraphs in the section
        full_content = f"# {title}\n\n" + "\n\n".join(content)

        # Create metadata
        metadata = {
            "doc_type": "section",
            "section_title": title,
            "section_level": level,
            "section_hierarchy": " > ".join(
                filter(None, self.current_section_hierarchy)
            ),
            "source_file": Path(source_file).name,
            "content_type": "text_section",
            "section_id": f"section_{self.section_counter}",
        }

        self.section_counter += 1

        return Document(page_content=full_content, metadata=metadata)

    def _process_table(
        self, table: Table, current_section: str, source_file: str
    ) -> List[Document]:
        """
        Process a table into multiple Document objects (one per row)
        """
        documents = []

        if not table.rows:
            return documents

        # Extract headers from first row
        headers = []
        first_row = table.rows[0]
        for cell in first_row.cells:
            headers.append(cell.text.strip())

        # Process each data row (skip header row)
        for row_idx, row in enumerate(table.rows[1:], 1):
            row_content = self._process_table_row(row, headers, row_idx)

            if row_content:
                # Create metadata for the table row
                metadata = {
                    "doc_type": "table_row",
                    "section_title": current_section,
                    "section_hierarchy": " > ".join(
                        filter(None, self.current_section_hierarchy)
                    ),
                    "source_file": Path(source_file).name,
                    "content_type": "table_data",
                    "table_headers": " | ".join(headers),
                    "row_number": row_idx,
                    "table_id": f"table_{current_section}_{row_idx}",
                }

                doc = Document(page_content=row_content, metadata=metadata)
                documents.append(doc)

        return documents

    def _process_table_row(
        self, row, headers: List[str], row_idx: int
    ) -> Optional[str]:
        """
        Process a single table row into content string
        """
        row_data = []

        for cell in row.cells:
            cell_text = cell.text.strip()
            row_data.append(cell_text)

        # Skip empty rows
        if not any(row_data):
            return None

        # Create structured content from the row
        content_parts = []

        for header, cell_value in zip(headers, row_data):
            if cell_value:  # Only include non-empty cells
                content_parts.append(f"{header}: {cell_value}")

        if not content_parts:
            return None

        # Create the final content
        row_content = f"Bảng dữ liệu - Hàng {row_idx}:\n" + "\n".join(content_parts)

        return row_content

    def get_document_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """
        Get statistics about the processed documents
        """
        stats = {
            "total_documents": len(documents),
            "sections": 0,
            "table_rows": 0,
            "doc_types": {},
            "sections_by_level": {},
        }

        for doc in documents:
            doc_type = doc.metadata.get("doc_type", "unknown")
            stats["doc_types"][doc_type] = stats["doc_types"].get(doc_type, 0) + 1

            if doc_type == "section":
                stats["sections"] += 1
                level = doc.metadata.get("section_level", 0)
                stats["sections_by_level"][level] = (
                    stats["sections_by_level"].get(level, 0) + 1
                )
            elif doc_type == "table_row":
                stats["table_rows"] += 1

        return stats


# Integration helper function
def integrate_with_knowledge_base(word_file_path: str) -> List[Document]:
    """
    Helper function to integrate with existing ViettelKnowledgeBase

    This can replace the manual CSV processing in the existing pipeline
    """
    processor = AutomatedWordProcessor()
    documents = processor.process_word_document(word_file_path)

    # Print stats
    stats = processor.get_document_stats(documents)
    print(f"[INFO] Document processing stats:")
    for key, value in stats.items():
        print(f"  {key}: {value}")

    return documents


# Example usage and testing
if __name__ == "__main__":
    # Test the processor
    processor = AutomatedWordProcessor()

    # Example file path (adjust as needed)
    test_file = "viettelpay_docs/raw/Nghiệp vụ.docx"

    try:
        documents = processor.process_word_document(test_file)

        # Show some example documents
        print(f"\n[INFO] Documents:")
        for i, doc in enumerate(documents):
            # if doc.metadata.get("doc_type") != "section":
            #     continue
            print(f"\nDocument {i+1}:")
            print(f"Type: {doc.metadata.get('doc_type')}")
            print(f"Section: {doc.metadata.get('section_title')}")
            print(f"Content preview: {doc.page_content[:150]}...")
            print(f"Metadata: {doc.metadata}")

        # Show stats
        stats = processor.get_document_stats(documents)
        print(f"\n[INFO] Processing statistics:")
        for key, value in stats.items():
            print(f"  {key}: {value}")

    except FileNotFoundError:
        print("[ERROR] Test file not found. Please adjust the file path.")
    except Exception as e:
        print(f"[ERROR] Error processing document: {e}")
